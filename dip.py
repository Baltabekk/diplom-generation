import asyncio
import logging
import os
import time
from aiogram import Bot, Dispatcher, types
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.filters.command import Command
from aiogram.types import FSInputFile
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import google.generativeai as genai
from google.api_core import retry, exceptions

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Установите ваши ключи API
TELEGRAM_TOKEN = '6790686493:AAGYsdk5DVLrmccPT87Li49P1RBEik60-48'
GENAI_API_KEY = 'AIzaSyDr_zy732Xybb1xZ1LEpEH31h6PjgnWInQ'
genai.configure(api_key=GENAI_API_KEY)
model = genai.GenerativeModel('gemini-pro')

# Инициализация бота и диспетчера
bot = Bot(token=TELEGRAM_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(storage=storage)

# Определение состояний
class GenerationStates(StatesGroup):
    GENERATING_SECTIONS = State()

# Определение пользовательской стратегии повторных попыток
custom_retry = retry.Retry(
    initial=1.0,
    maximum=60.0,
    multiplier=2.0,
    predicate=retry.if_exception_type(
        exceptions.DeadlineExceeded,
        exceptions.ServiceUnavailable,
    ),
)

@dp.message(Command("start"))
async def send_welcome(message: types.Message):
    logger.info(f"Пользователь {message.from_user.id} начал взаимодействие")
    await message.reply("Привет! Я бот, который поможет вам создать дипломную работу. Напишите /help для списка команд.")

@dp.message(Command("help"))
async def send_help(message: types.Message):
    logger.info(f"Пользователь {message.from_user.id} запросил помощь")
    await message.reply("Доступные команды:\n/start - Начать взаимодействие\n/generate [тема] - Создать содержание и текст\n/help - Справка")

@dp.message(Command("generate"))
async def generate_content(message: types.Message, state: FSMContext):
    topic = message.text[len('/generate '):].strip()
    if not topic:
        logger.warning(f"Пользователь {message.from_user.id} не указал тему")
        await message.reply("Пожалуйста, укажите тему после команды /generate.")
        return

    logger.info(f"Пользователь {message.from_user.id} запросил генерацию на тему: {topic}")
    await message.reply("Генерация содержания...")
    prompt = f"Создайте подробное содержание для дипломной работы на тему '{topic}'. Включите введение, 5-7 основных разделов с 3-4 подразделами каждый, и заключение."

    try:
        response = custom_retry(model.generate_content)(prompt)
        content = response.text.strip()
        logger.info(f"Содержание сгенерировано для пользователя {message.from_user.id}")
        await message.reply(f"Содержание:\n{content}")

        sections = ["Введение"] + [section.strip() for section in content.split('\n') if section.strip()] + ["Заключение"]
        await state.update_data(topic=topic, sections=sections, results=[])
        
        # Отправляем сообщение о начале генерации и примерном времени
        status_message = await message.reply("Начинаю генерацию документа. Это займет примерно 13 минут.\nПрогресс: 0%")
        
        await state.set_state(GenerationStates.GENERATING_SECTIONS)
        await generate_sections(message, state, status_message)
    except Exception as e:
        logger.error(f"Ошибка при генерации содержания: {str(e)}", exc_info=True)
        await message.reply("Произошла ошибка при генерации содержания. Пожалуйста, попробуйте еще раз.")

async def generate_sections(message: types.Message, state: FSMContext, status_message: types.Message):
    data = await state.get_data()
    sections = data.get('sections', [])
    results = data.get('results', [])
    topic = data.get('topic')

    total_sections = len(sections)
    start_time = time.time()

    for i, section in enumerate(sections):
        logger.info(f"Генерация раздела '{section}' для пользователя {message.from_user.id}")
        
        if section == "Введение":
            prompt = f"Напишите подробное введение (около 500 слов) для дипломной работы на тему '{topic}'. Включите актуальность темы, цели и задачи исследования."
        elif section == "Заключение":
            prompt = f"Напишите подробное заключение (около 500 слов) для дипломной работы на тему '{topic}'. Подведите итоги исследования, сформулируйте основные выводы."
        else:
            prompt = f"Напишите подробный текст (около 1000-1500 слов) для раздела '{section}' дипломной работы на тему '{topic}'. Включите теоретическую базу, анализ и практические аспекты."

        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = custom_retry(model.generate_content)(prompt)
                section_text = response.text.strip()
                results.append((section, section_text))
                logger.info(f"Раздел '{section}' сгенерирован для пользователя {message.from_user.id}")
                break
            except exceptions.DeadlineExceeded as e:
                if attempt < max_retries - 1:
                    wait_time = (attempt + 1) * 5
                    logger.warning(f"Timeout error for section '{section}'. Retrying in {wait_time} seconds...")
                    await asyncio.sleep(wait_time)
                else:
                    logger.error(f"Failed to generate section '{section}' after {max_retries} attempts: {str(e)}")
            except Exception as e:
                logger.error(f"Ошибка при генерации раздела '{section}': {str(e)}", exc_info=True)
                break

        await update_progress(status_message, start_time, i + 1, total_sections)

    await state.update_data(results=results)
    await finalize_document(message, state, status_message)

async def update_progress(status_message: types.Message, start_time: float, completed_sections: int, total_sections: int):
    progress = int((completed_sections / total_sections) * 100)
    elapsed_time = int(time.time() - start_time)
    await status_message.edit_text(f"Генерация документа.\nПрошло времени: {elapsed_time // 60} мин {elapsed_time % 60} сек\nПрогресс: {progress}%")

async def finalize_document(message: types.Message, state: FSMContext, status_message: types.Message):
    data = await state.get_data()
    topic = data.get('topic')
    results = data.get('results', [])

    if not results:
        logger.error(f"Ошибка при создании документа для пользователя {message.from_user.id}")
        await message.reply("Ошибка при создании документа.")
        await state.clear()
        return

    logger.info(f"Создание итогового документа для пользователя {message.from_user.id}")
    await status_message.edit_text("Создание итогового документа... Прогресс: 100%")
    
    doc = Document()
    
    # Создаем стиль для оглавления
    styles = doc.styles
    toc_style = styles.add_style('TOC', WD_STYLE_TYPE.PARAGRAPH)
    toc_style.font.size = Pt(12)
    toc_style.font.name = 'Times New Roman'

    doc.add_heading(f"Дипломная работа на тему: {topic}", level=0)
    
    # Добавление оглавления
    doc.add_paragraph("Оглавление")
    for section, _ in results:
        doc.add_paragraph(section, style='TOC')
    
    doc.add_page_break()

    for section, text in results:
        doc.add_heading(section, level=1)
        paragraphs = text.split('\n\n')  # Разделяем текст на параграфы
        for para in paragraphs:
            doc.add_paragraph(para)
        doc.add_page_break()  # Добавляем разрыв страницы после каждого раздела

    file_path = f"{message.from_user.id}_diploma.docx"
    doc.save(file_path)
    logger.info(f"Документ сохранен: {file_path}")

    input_file = FSInputFile(file_path)
    logger.info(f"Отправка документа пользователю {message.from_user.id}")
    await message.reply_document(input_file, caption="Ваша дипломная работа готова!")

    os.remove(file_path)
    logger.info(f"Документ удален с сервера: {file_path}")
    await message.reply("Документ успешно отправлен и удален с сервера.")
    await state.clear()

async def main():
    logger.info("Бот запущен")
    await dp.start_polling(bot)

if __name__ == '__main__':
    asyncio.run(main())
