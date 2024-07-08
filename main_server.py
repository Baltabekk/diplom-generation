import asyncio
import logging
import os
from aiogram import Bot, Dispatcher, types
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.filters import Command
from aiogram.types import FSInputFile
from docx import Document
from docx.shared import Pt
from cachetools import TTLCache
import requests

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)
TELEGRAM_TOKEN = os.getenv('TELEGRAM_TOKEN')
ADMIN_ID = 123456789 

# Инициализация бота и диспетчера
bot = Bot(token=TELEGRAM_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(storage=storage)

# Кэш для хранения сгенерированного контента
cache = TTLCache(maxsize=100, ttl=3600)

# Определение состояний
class GenerationStates(StatesGroup):
    GENERATING_SECTIONS = State()

@dp.message(Command("start"))
async def send_welcome(message: types.Message):
    logger.info(f"Пользователь {message.from_user.id} начал взаимодействие")
    await message.reply("Привет! Я бот, который поможет вам создать дипломную работу. Напишите /help для списка команд.")

@dp.message(Command("help"))
async def send_help(message: types.Message):
    logger.info(f"Пользователь {message.from_user.id} запросил помощь")
    await message.reply(
        "Доступные команды:\n"
        "/start - Начать взаимодействие\n"
        "/generate [тема] - Создать содержание и текст\n"
        "/example - Посмотреть пример\n"
        "/feedback - Оставить отзыв\n"
        "/contact_admin - Связаться с администратором\n"
        "/help - Справка"
    )

@dp.message(Command("example"))
async def send_example(message: types.Message):
    logger.info(f"Пользователь {message.from_user.id} запросил пример")
    await message.reply("Пример дипломной работы:\n\nВведение\n...\nРаздел 1\n...\nРаздел 2\n...\nЗаключение\n...")

@dp.message(Command("feedback"))
async def send_feedback(message: types.Message):
    logger.info(f"Пользователь {message.from_user.id} запросил оставить отзыв")
    await message.reply("Пожалуйста, отправьте ваш отзыв сообщением. Я передам его администратору.")

@dp.message(Command("contact_admin"))
async def contact_admin(message: types.Message):
    logger.info(f"Пользователь {message.from_user.id} запросил связаться с администратором")
    await message.reply("Если у вас есть вопросы, вы можете связаться с администратором. Напишите /message_admin [ваше сообщение].")

@dp.message(Command("message_admin"))
async def message_admin(message: types.Message):
    admin_message = message.text[len('/message_admin '):].strip()
    if not admin_message:
        await message.reply("Пожалуйста, укажите сообщение после команды /message_admin.")
        return
    logger.info(f"Пользователь {message.from_user.id} отправил сообщение администратору: {admin_message}")
    await bot.send_message(ADMIN_ID, f"Сообщение от пользователя {message.from_user.username} ({message.from_user.id}):\n\n{admin_message}")
    await message.reply("Ваше сообщение было отправлено администратору.")

@dp.message(Command("broadcast"))
async def broadcast_message(message: types.Message):
    if message.from_user.id != ADMIN_ID:
        await message.reply("У вас нет прав для выполнения этой команды.")
        return
    broadcast_message = message.text[len('/broadcast '):].strip()
    if not broadcast_message:
        await message.reply("Пожалуйста, укажите сообщение после команды /broadcast.")
        return
    logger.info(f"Администратор отправляет сообщение всем пользователям: {broadcast_message}")
    for user_id in users:
        try:
            await bot.send_message(user_id, broadcast_message)
        except Exception as e:
            logger.error(f"Ошибка при отправке сообщения пользователю {user_id}: {str(e)}", exc_info=True)
    await message.reply("Сообщение было отправлено всем пользователям.")

@dp.message(Command("list_users"))
async def list_users(message: types.Message):
    if message.from_user.id != ADMIN_ID:
        await message.reply("У вас нет прав для выполнения этой команды.")
        return
    logger.info(f"Администратор запросил список пользователей")
    user_list = "Список пользователей:\n\n"
    for user_id, user_info in users.items():
        user_list += f"ID: {user_id}, Username: {user_info.get('username')}, Имя: {user_info.get('first_name')} {user_info.get('last_name')}, Дата: {user_info.get('date')}\n"
    await message.reply(user_list)

async def generate_content_with_cache(prompt, user_id):
    cache_key = f"{user_id}:{prompt}"
    if cache_key in cache:
        return cache[cache_key]

    # Отправляем запрос рабочему серверу
    worker_url = await get_worker_url()
    response = requests.post(worker_url, json={'prompt': prompt, 'user_id': user_id})
    if response.status_code == 200:
        content = response.json()['content']
        cache[cache_key] = content
        return content
    else:
        raise Exception(f"Ошибка при генерации контента: {response.text}")

@dp.message(Command("generate"))
async def generate_content(message: types.Message, state: FSMContext):
    topic = message.text[len('/generate '):].strip()
    if not topic:
        logger.warning(f"Пользователь {message.from_user.id} не указал тему")
        await message.reply("Пожалуйста, укажите тему после команды /generate.")
        return

    logger.info(f"Пользователь {message.from_user.id} запросил генерацию на тему: {topic}")
    await message.reply("Генерация содержания...")
    prompt = f"Создайте подробное содержание для дипломной работы на тему '{topic}'. Включите введение, 5 основных разделов с 3 подразделами каждый, обязательно номер для каждого раздела и подраздела, и заключение."

    try:
        content = await generate_content_with_cache(prompt, message.from_user.id)
        logger.info(f"Содержание сгенерировано для пользователя {message.from_user.id}")
        await message.reply(f"Содержание:\n{content}")

        sections = parse_content(content)
        await state.update_data(topic=topic, sections=sections)
        
        status_message = await message.reply("Начинаю генерацию документа. Это займет некоторое время.\nПрогресс: 0%")
        
        await state.set_state(GenerationStates.GENERATING_SECTIONS)
        await generate_sections(message, state, status_message)
    except Exception as e:
        logger.error(f"Ошибка при генерации содержания: {str(e)}", exc_info=True)
        await message.reply(f"Произошла ошибка при генерации содержания. Пожалуйста, попробуйте еще раз или обратитесь к администратору.")

def parse_content(content):
    sections = []
    for line in content.split('\n'):
        if line.strip():
            level = 0
            while line.startswith(' '):
                level += 1
                line = line[1:].strip()
            sections.append((level, line))
    return sections

async def generate_section(topic, section, user_id):
    prompt = f"Создайте подробный текст для раздела '{section}' дипломной работы на тему '{topic}'."
    # Отправляем запрос рабочему серверу
    worker_url = await get_worker_url()
    response = requests.post(worker_url, json={'prompt': prompt, 'user_id': user_id})
    if response.status_code == 200:
        return response.json()['content']
    else:
        raise Exception(f"Ошибка при генерации контента: {response.text}")

async def generate_sections(message, state, status_message):
    data = await state.get_data()
    topic = data['topic']
    sections = data['sections']
    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    for i, (level, section) in enumerate(sections):
        if level == 0:
            doc.add_heading(section, level=1)
        elif level == 1:
            doc.add_heading(section, level=2)
        else:
            doc.add_heading(section, level=3)
        
        text = await generate_section(topic, section, message.from_user.id)
        doc.add_paragraph(text)
        
        progress = int((i + 1) / len(sections) * 100)
        await status_message.edit_text(f"Прогресс: {progress}%")
    
    filename = f"{topic.replace(' ', '_')}_diploma.docx"
    filepath = os.path.join(os.getcwd(), filename)
    doc.save(filepath)
    
    await message.reply_document(FSInputFile(filepath))
    os.remove(filepath)
    await state.clear()
    await status_message.delete()

async def get_worker_url():
    # Здесь вы должны реализовать логику выбора рабочего сервера
    # Например, можно использовать очередь или другой механизм балансировки нагрузки
    worker_urls = [
        "http://worker1_server:8000/generate",
        "http://worker2_server:8000/generate",
        "http://worker3_server:8000/generate",
        "http://worker4_server:8000/generate",
        "http://worker5_server:8000/generate"
    ]
    for url in worker_urls:
        response = requests.get(url + "/status")
        if response.status_code == 200 and response.json().get("status") == "ready":
            return url
    raise Exception("Все рабочие сервера заняты. Попробуйте позже.")

if __name__ == "__main__":
    dp.run_polling(bot)
