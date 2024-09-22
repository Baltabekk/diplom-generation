import asyncio
import logging
from datetime import datetime, timedelta
from aiogram import Bot, Dispatcher, types, F
from aiogram.enums import ParseMode
from aiogram.filters import Command
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import FSInputFile, ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardMarkup, InlineKeyboardButton
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import google.generativeai as genai
from google.api_core import retry, exceptions
from cachetools import TTLCache
import random
import json
import os

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Константы
ADMIN_ID = 1419048544
DAILY_LIMIT = 3
TELEGRAM_TOKEN = '7414905635:AAHBlef17Zjo0x13nrTCV0X410fiyY1TOKQ'
GENAI_API_KEYS = [
    'AIzaSyAnxCZjiQqWgkX-tLIR7K9OROqAVVDKNBw'
]# Пути к файлам данных
USER_DATA_FILE = 'user_data.json'
FEEDBACK_FILE = 'feedback.json'

# Инициализация бота и хранилища
bot = Bot(token=TELEGRAM_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(storage=storage)

# Кэш для хранения сгенерированного контента
cache = TTLCache(maxsize=100, ttl=3600)

# Определение состояний
class GenerationStates(StatesGroup):
    SELECT_DOCUMENT_SIZE = State()
    WAITING_FOR_TOPIC = State()
    GENERATING_CONTENT = State()
    LEAVE_FEEDBACK = State()
    SEND_MESSAGE_TO_ALL = State()

# Пользовательская стратегия повторных попыток
custom_retry = retry.Retry(
    initial=1.0,
    maximum=60.0,
    multiplier=2.0,
    predicate=retry.if_exception_type(
        exceptions.DeadlineExceeded, exceptions.ServiceUnavailable
    )
)

# Функции для работы с данными
def load_data(file_path, default=None):
    if os.path.exists(file_path):
        with open(file_path, 'r') as file:
            return json.load(file)
    return default if default is not None else {}

def save_data(data, file_path):
    with open(file_path, 'w') as file:
        json.dump(data, file, indent=4, default=str)

# Загрузка данных при запуске бота
user_data = load_data(USER_DATA_FILE, {})
feedback_storage = load_data(FEEDBACK_FILE, [])

def get_user_data(user_id):
    user_id = str(user_id)
    if user_id not in user_data:
        user_data[user_id] = {
            'requests': 0,
            'last_reset': str(datetime.now()),
            'is_admin': user_id == str(ADMIN_ID),
            'start_date': str(datetime.now()),
            'username': None,
            'feedback': [],
            'requested_topic': None,
            'referral_link': f"https://t.me/gendiplom_bot?start={user_id}",
            'referred_by': None,
            'referral_count': 0,
            'bonus_requests': 0
        }
    return user_data[user_id]

async def process_referral(new_user_id, referrer_id):
    new_user_data = get_user_data(new_user_id)

    # Проверка, был ли пользователь уже приглашен
    if new_user_data['referred_by']:
        return False, "Вы уже были приглашены другим пользователем."

    # Проверка, не использует ли пользователь свою собственную реферальную ссылку
    if str(new_user_id) == str(referrer_id):
        return False, "Вы не можете использовать свою собственную реферальную ссылку."

    # Получаем данные реферера
    referrer_data = get_user_data(referrer_id)

    # Обновляем данные пользователя и реферера
    new_user_data['referred_by'] = referrer_id
    referrer_data['referral_count'] += 1
    referrer_data['bonus_requests'] += 2
    new_user_data['bonus_requests'] += 3

    # Сохраняем обновленные данные
    save_data(new_user_data, USER_DATA_FILE)  # Обновите данные нового пользователя

    # Отправляем уведомление рефереру
    try:
        await bot.send_message(
            int(referrer_id),
            f"Поздравляем! Вы пригласили нового пользователя и получили 2 дополнительных запроса. Ваш текущий бонус: {referrer_data['bonus_requests']} запросов."
        )
    except Exception as e:
        logger.error(f"Не удалось отправить уведомление пользователю {referrer_id}: {str(e)}")
    
    return True, "Вы успешно присоединились по реферальной ссылке! Вы получили 3 дополнительных запроса."

def check_and_update_quota(user_id, is_document_generation=False):
    user_id = str(user_id)
    data = get_user_data(user_id)
    if data['is_admin']:
        return True, "Админ имеет неограниченный доступ."

    now = datetime.now()
    last_reset = datetime.fromisoformat(data['last_reset'])
    if (now - last_reset).days >= 1:
        data['requests'] = 0
        data['last_reset'] = str(now)

    total_requests = DAILY_LIMIT + data['bonus_requests']

    if is_document_generation and data['requests'] >= total_requests:
        next_reset = last_reset + timedelta(days=1)
        return False, f"Достигнут дневной лимит. Следующее обновление: {next_reset.strftime('%Y-%m-%d %H:%M:%S')}."

    if is_document_generation:
        data['requests'] += 1
        if data['requests'] > DAILY_LIMIT and data['bonus_requests'] > 0:
            data['bonus_requests'] -= 1
        remaining = total_requests - data['requests']
        save_data(user_data, USER_DATA_FILE)
        return True, f"Запрос принят. Осталось запросов сегодня: {remaining}."
    
    return True, f"Осталось запросов сегодня: {total_requests - data['requests']}."

def get_remaining_requests(user_id):
    data = get_user_data(str(user_id))
    if data['is_admin']:
        return "Неограниченно (админ)"
    total_requests = DAILY_LIMIT + data['bonus_requests']
    remaining = total_requests - data['requests']
    return str(remaining)

def get_random_api_key():
    return random.choice(GENAI_API_KEYS)

def initialize_model(api_key):
    genai.configure(api_key=api_key)
    return genai.GenerativeModel('gemini-1.5-flash')

current_model = initialize_model(get_random_api_key())

# Функция для создания главной клавиатуры
def get_main_menu_keyboard():
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="Создать дипломную работу")],
            [KeyboardButton(text="Оставить отзыв"), KeyboardButton(text="Посмотреть отзывы")],
            [KeyboardButton(text="FAQ"), KeyboardButton(text="О нас")],
            [KeyboardButton(text="Моя реферальная ссылка")]
        ],
        resize_keyboard=True
    )

@dp.message(Command("start"))
async def start_command(message: types.Message):
    user_id = message.from_user.id
    user_data = get_user_data(user_id)
    
    # Получение аргументов из текста сообщения
    args = message.text.split()[1:]  # Все аргументы после команды
    if args:
        referrer_id = args[0]  # Первый аргумент
        success, referral_message = await process_referral(user_id, referrer_id)
        if success:
            await message.reply(referral_message)
    
    welcome_text = f"Привет, {message.from_user.first_name}!\nДобро пожаловать в наш бот."
    await message.reply(welcome_text, reply_markup=get_main_menu_keyboard())


@dp.message(Command("help"))
async def send_help(message: types.Message):
    logger.info(f"Пользователь {message.from_user.id} запросил помощь")
    help_text = (
        "Доступные команды:\n"
        "/start - Показать главное меню\n"
        "/generate - Создать содержание и текст\n"
        "/quota - Проверить оставшиеся запросы\n"
        "/leave_feedback - Оставить отзыв\n"
        "/view_feedback - Посмотреть отзывы\n"
        "/contact_admins - Связаться с админами\n"
        "/admin_menu - Меню администратора (только для админов)\n"
        "/faq - Часто задаваемые вопросы\n"
        "/about_us - О нашей компании\n"
        "/my_referral - Ваша реферальная ссылка"
    )
    await message.reply(help_text, reply_markup=get_main_menu_keyboard())

@dp.message(Command("generate"))
@dp.message(F.text == "Создать дипломную работу")
async def generate_command(message: types.Message, state: FSMContext):
    markup = ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="40 страниц"), KeyboardButton(text="60 страниц"), KeyboardButton(text="100 страниц")],
            [KeyboardButton(text="Отменить генерацию")],
            [KeyboardButton(text="Главное меню")]
        ],
        resize_keyboard=True,
        one_time_keyboard=True
    )
    await message.reply("Выберите размер документа или отмените генерацию:", reply_markup=markup)
    await state.set_state(GenerationStates.SELECT_DOCUMENT_SIZE)

@dp.message(GenerationStates.SELECT_DOCUMENT_SIZE, F.text == "Отменить генерацию")
async def cancel_generation(message: types.Message, state: FSMContext):
    await state.clear()
    await message.reply("Генерация отменена.", reply_markup=get_main_menu_keyboard())

@dp.message(F.text == "Главное меню")
async def main_menu(message: types.Message, state: FSMContext):
    await state.clear()
    await message.reply("Вы вернулись в главное меню.", reply_markup=get_main_menu_keyboard())

@dp.message(GenerationStates.SELECT_DOCUMENT_SIZE)
async def process_document_size(message: types.Message, state: FSMContext):
    size = message.text.lower()
    word_count = 0
    
    if "40" in size:
        word_count = 500
    elif "60" in size:
        word_count = 1000
    elif "100" in size:
        word_count = 2000
    else:
        await message.reply("Пожалуйста, выберите размер документа из предложенных вариантов.")
        return

    await state.update_data(word_count=word_count)
    await message.reply("Отлично! Теперь введите тему вашей дипломной работы.")
    await state.set_state(GenerationStates.WAITING_FOR_TOPIC)

@dp.message(GenerationStates.WAITING_FOR_TOPIC)
async def receive_topic(message: types.Message, state: FSMContext):
    topic = message.text.strip()
    await state.update_data(topic=topic)
    await message.reply(f"Тема принята: {topic}\nТеперь я начну генерацию содержания.")
    await state.set_state(GenerationStates.GENERATING_CONTENT)
    await generate_content(message, state)

@dp.message(GenerationStates.GENERATING_CONTENT)
async def ignore_messages_during_generation(message: types.Message):
    await message.reply("Пожалуйста, подождите. Идет генерация содержания. Вы получите уведомление, когда документ будет готов.")

@dp.message(F.text == "Моя реферальная ссылка")
async def my_referral(message: types.Message):
    user_id = str(message.from_user.id)
    user_data = get_user_data(user_id)
    referral_link = f"https://t.me/gendiplom_bot?start={user_id}"
    referral_count = user_data['referral_count']
    bonus_requests = user_data['bonus_requests']
    
    response = (
        f"Ваша реферальная ссылка: {referral_link}\n\n"
        f"Количество приглашенных пользователей: {referral_count}\n"
        f"Ваш текущий бонус: {bonus_requests} дополнительных запросов\n\n"
        "Отправьте эту ссылку друзьям. За каждого нового пользователя вы получите 2 дополнительных запроса, а ваш друг - 3 запроса!"
    )
    
    await message.reply(response, reply_markup=get_main_menu_keyboard())


async def generate_content(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    data = await state.get_data()
    topic = data.get('topic', '')
    doc_size = data.get('word_count', 300)

    logger.info(f"Пользователь {user_id} указал тему: {topic}")

    user_data = get_user_data(user_id)
    user_data['requested_topic'] = topic
    can_proceed, quota_message = check_and_update_quota(user_id, is_document_generation=True)

    if not can_proceed:
        await message.reply(quota_message, reply_markup=get_main_menu_keyboard())
        return

    await message.reply(f"{quota_message}\nГенерация содержания...")

    section_count = 4 if doc_size == 500 else 6 if doc_size == 1000 else 7
    subsection_count = (2, 3) if doc_size == 500 else (3, 4) if doc_size == 1000 else (4, 5)

    # Улучшенный промпт для генерации содержания
    prompt = (f"Создайте подробное содержание для дипломной работы на тему '{topic}'. "
              f"Необходимо включить введение, {section_count} основных разделов, каждый из которых, не пишите ничего лишнего каждый ваш абзац считается содержанием для документа , ничего не спрашивай при создании содержание"
              f"содержит от {subsection_count[0]} до {subsection_count[1]} подразделов. "
              f"Закончите заключением. Избегайте повторений разделов и подразделов, "
              f"используйте академический стиль.")

    async def fetch_with_retries(prompt, retries=3):
        for attempt in range(retries):
            try:
                response = await asyncio.wait_for(asyncio.to_thread(current_model.generate_content, prompt), timeout=60)
                return response.text
            except Exception as e:
                logger.error(f"Попытка {attempt + 1} не удалась: {str(e)}")
                if attempt < retries - 1:
                    await asyncio.sleep(2)  # Подождите перед повтором
        raise Exception("Не удалось получить ответ после нескольких попыток.")

    try:
        content = await fetch_with_retries(prompt)
        logger.info(f"Содержание сгенерировано для пользователя {user_id}")

        # Функция для отправки длинного сообщения
        async def send_long_message(chat_id, text):
            for chunk in [text[i:i + 4096] for i in range(0, len(text), 4096)]:
                await message.reply(chunk)

        await send_long_message(user_id, f"Содержание:\n{content}")

        sections = ["Введение"] + [section.strip() for section in content.split('\n') if section.strip() and not section.startswith('-')] + ["Заключение"]
        await state.update_data(sections=sections, words_per_section=doc_size)

        status_message = await message.reply("Начинаю генерацию документа. Это займет некоторое время.\nПрогресс: 0%")

        await generate_sections(message, state, status_message)

    except asyncio.TimeoutError:
        error_message = "Генерация содержания заняла слишком много времени. Попробуйте еще раз."
        logger.error(error_message)
        await message.reply(error_message, reply_markup=get_main_menu_keyboard())

    except Exception as e:
        error_message = f"Ошибка при генерации содержания: {str(e)}"
        logger.error(error_message, exc_info=True)
        await message.reply(error_message, reply_markup=get_main_menu_keyboard())


async def generate_section(topic: str, section: str, words_per_section: int, previous_content: str) -> tuple:
    # Промпт для генерации раздела с учётом уже сгенерированных данных
    prompt = (f"Напишите раздел '{section}' для дипломной работы на тему '{topic}'. "
              f"Уже сгенерированные разделы:\n{previous_content}\n"
              f"Не повторяйте информацию, уже указанную в других разделах. "
              f"Используйте эти данные для создания непротиворечивого заключения. "
              f"Объем текста должен быть  {words_per_section} слов.")
    
    try:
        response = await asyncio.wait_for(asyncio.to_thread(current_model.generate_content, prompt), timeout=60)
        return section, response.text
    
    except asyncio.TimeoutError:
        logger.error(f"Тайм-аут при генерации раздела '{section}'")
        return section, f"Генерация раздела '{section}' заняла слишком много времени."
    
    except Exception as e:
        logger.error(f"Ошибка при генерации раздела '{section}': {str(e)}")
        return section, f"Не удалось сгенерировать содержание для раздела '{section}'"

async def generate_sections(message: types.Message, state: FSMContext, status_message: types.Message):
    data = await state.get_data()
    topic = data['topic']
    sections = data['sections']
    words_per_section = data['words_per_section']
    
    results = []
    total_sections = len(sections)
    start_time = asyncio.get_event_loop().time()
    
    previous_content = ""  # Храним сгенерированные данные
    
    for i, section in enumerate(sections):
        section_content = await generate_section(topic, section, words_per_section, previous_content)
        previous_content += f"\n\n{section_content[1]}"  # Добавляем сгенерированный раздел к предыдущим данным
        results.append(section_content)
        
        await update_progress(status_message, start_time, i + 1, total_sections)
    
    await state.update_data(results=results)
    await finalize_document(message, state, status_message)


async def update_progress(status_message, start_time, current, total):
    elapsed_time = asyncio.get_event_loop().time() - start_time
    progress = (current / total) * 100
    estimated_total_time = elapsed_time / (current / total)
    estimated_remaining_time = estimated_total_time - elapsed_time
    
    await status_message.edit_text(
        f"Генерация документа... Прогресс: {progress:.2f}%\n"
        f"Прошло времени: {elapsed_time:.2f} сек.\n"
        f"Осталось примерно: {estimated_remaining_time:.2f} сек."
    )

async def finalize_document(message: types.Message, state: FSMContext, status_message: types.Message):
    data = await state.get_data()
    topic = data.get('topic')
    results = data.get('results', [])

    if not topic or not results:
        logger.error(f"Ошибка: тема или результаты отсутствуют для пользователя {message.from_user.id}")
        await message.reply("Ошибка при создании документа: отсутствуют тема или содержание.", reply_markup=get_main_menu_keyboard())
        await state.clear()
        return

    if all(not content for _, content in results):
        logger.error(f"Ошибка: все разделы пустые для пользователя {message.from_user.id}")
        await message.reply("Ошибка при создании документа: все разделы пустые.", reply_markup=get_main_menu_keyboard())
        await state.clear()
        return

    logger.info(f"Создание итогового документа для пользователя {message.from_user.id}")
    await status_message.edit_text("Создание итогового документа... Прогресс: 100%")

    doc = Document()
    styles = doc.styles
    toc_style = styles.add_style('TOC', WD_STYLE_TYPE.PARAGRAPH)
    toc_style.font.size = Pt(12)
    toc_style.font.name = 'Times New Roman'

    doc.add_heading(f"Дипломная работа на тему: {topic}", level=0)

    doc.add_paragraph("Оглавление")
    for section, _ in results:
        doc.add_paragraph(section, style='TOC')

    doc.add_page_break()

    for section, text in results:
        doc.add_heading(section, level=1)
        paragraphs = text.split('\n\n')

        for para in paragraphs:
            new_para = doc.add_paragraph()
            if para.startswith('**') and para.endswith('**'):
                clean_text = para[2:-2].strip()
                run = new_para.add_run(clean_text)
                run.bold = True
            else:
                parts = para.split('**')
                for i, part in enumerate(parts):
                    part = part.strip()
                    if i % 2 == 0:
                        if part.startswith('* '):
                            part = '• ' + part[2:].capitalize()
                        new_para.add_run(part)
                    else:
                        run = new_para.add_run(part.strip())
                        run.bold = True

    file_path = f"{message.from_user.id}_{topic.replace(' ', '_')}.docx"
    doc.save(file_path)
    logger.info(f"Документ сохранен: {file_path}")

    input_file = FSInputFile(file_path)
    logger.info(f"Отправка документа пользователю {message.from_user.id}")
    await message.reply_document(
        input_file, 
        caption=f"Ваш документ на тему '{topic}' готов!", 
        reply_markup=get_main_menu_keyboard()
    )

    await state.clear()


@dp.message(Command("quota"))
async def check_quota(message: types.Message):
    user_id = message.from_user.id
    remaining = get_remaining_requests(user_id)
    await message.reply(f"Оставшиеся запросы на сегодня: {remaining}", reply_markup=get_main_menu_keyboard())

@dp.message(Command("leave_feedback"))
@dp.message(F.text == "Оставить отзыв")
async def leave_feedback(message: types.Message, state: FSMContext):
    await message.reply("Пожалуйста, напишите ваш отзыв.")
    await state.set_state(GenerationStates.LEAVE_FEEDBACK)

@dp.message(GenerationStates.LEAVE_FEEDBACK)
async def process_feedback(message: types.Message, state: FSMContext):
    feedback = message.text.strip()
    feedback_storage.append(feedback)
    user_id = str(message.from_user.id)
    user_data[user_id]['feedback'].append(feedback)
    save_data(feedback_storage, FEEDBACK_FILE)
    save_data(user_data, USER_DATA_FILE)
    await message.reply("Спасибо за ваш отзыв! Он был успешно сохранен.", reply_markup=get_main_menu_keyboard())
    await state.clear()

@dp.message(Command("view_feedback"))
@dp.message(F.text == "Посмотреть отзывы")
async def view_feedback(message: types.Message):
    if not feedback_storage:
        await message.reply("Отзывов пока нет.", reply_markup=get_main_menu_keyboard())
    else:
        feedback_list = "\n\n".join([f"Отзыв {i+1}: {feedback}" for i, feedback in enumerate(feedback_storage[-5:])])
        await message.reply(f"Последние 5 отзывов:\n\n{feedback_list}", reply_markup=get_main_menu_keyboard())

@dp.message(Command("contact_admins"))
async def contact_admins(message: types.Message):
    await message.reply("Для связи с администраторами, пожалуйста, напишите на: @baltabek_kk", reply_markup=get_main_menu_keyboard())

@dp.message(Command("admin_menu"))
async def admin_menu(message: types.Message):
    if str(message.from_user.id) != str(ADMIN_ID):
        await message.reply("У вас нет прав для доступа к этому меню.", reply_markup=get_main_menu_keyboard())
        return

    markup = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Просмотр участников", callback_data="view_participants")],
        [InlineKeyboardButton(text="Просмотр отзывов", callback_data="view_all_feedback")],
        [InlineKeyboardButton(text="Отправить сообщение всем", callback_data="send_message_to_all")]
    ])
    await message.reply("Меню администратора:", reply_markup=markup)

@dp.callback_query(F.data == "view_participants")
async def view_participants(callback_query: types.CallbackQuery):
    if str(callback_query.from_user.id) != str(ADMIN_ID):
        await callback_query.answer("У вас нет прав для выполнения этой команды.")
        return
    
    participants = "\n".join([f"ID: {user_id}, Username: {data.get('username', 'Не указан')}" for user_id, data in user_data.items()])
    await callback_query.message.reply(f"Список участников:\n{participants}")

@dp.callback_query(F.data == "view_all_feedback")
async def view_all_feedback(callback_query: types.CallbackQuery):
    if str(callback_query.from_user.id) != str(ADMIN_ID):
        await callback_query.answer("У вас нет прав для выполнения этой команды.")
        return
    
    if not feedback_storage:
        await callback_query.message.reply("Отзывов пока нет.")
    else:
        feedback_list = "\n\n".join([f"Отзыв {i+1}: {feedback}" for i, feedback in enumerate(feedback_storage)])
        await callback_query.message.reply(f"Все отзывы:\n\n{feedback_list}")

@dp.callback_query(F.data == "send_message_to_all")
async def send_message_to_all(callback_query: types.CallbackQuery, state: FSMContext):
    if str(callback_query.from_user.id) != str(ADMIN_ID):
        await callback_query.answer("У вас нет прав для выполнения этой команды.")
        return
    
    await callback_query.message.reply("Введите сообщение для всех пользователей.")
    await state.set_state(GenerationStates.SEND_MESSAGE_TO_ALL)

@dp.message(GenerationStates.SEND_MESSAGE_TO_ALL)
async def process_send_message_to_all(message: types.Message, state: FSMContext):
    if str(message.from_user.id) != str(ADMIN_ID):
        await message.reply("У вас нет прав для выполнения этой команды.")
        await state.clear()
        return
    
    message_text = message.text.strip()
    for user_id in user_data.keys():
        try:
            await bot.send_message(int(user_id), message_text)
        except Exception as e:
            logger.error(f"Не удалось отправить сообщение пользователю {user_id}: {str(e)}")
    await message.reply("Сообщение отправлено всем пользователям.")
    await state.clear()

@dp.message(Command("faq"))
@dp.message(F.text == "FAQ")
async def faq(message: types.Message):
    faq_text = (
        "Часто задаваемые вопросы:\n\n"
        "1. Как использовать бота?\n"
        "   Отправьте команду /generate, выберите размер документа и введите тему.\n\n"
        "2. Сколько запросов я могу сделать в день?\n"
        "   Обычные пользователи могут сделать 3 запроса в день.\n\n"
        "3. Как оставить отзыв?\n"
        "   Используйте команду /leave_feedback или кнопку 'Оставить отзыв'.\n\n"
        "4. Как связаться с администраторами?\n"
        "   Используйте команду /contact_admins.\n\n"
        "5. Как работает реферальная система?\n"
        "   За каждого приглашенного пользователя вы получаете 2 дополнительных запроса."
    )
    await message.reply(faq_text, reply_markup=get_main_menu_keyboard())

@dp.message(Command("about_us"))
@dp.message(F.text == "О нас")
async def about_us(message: types.Message):
    about_text = (
        "О нашей компании:\n\n"
        "Мы - команда энтузиастов, разрабатывающая инновационные решения "
        "для помощи студентам в их академической деятельности. "
        "Наш бот использует передовые технологии искусственного интеллекта "
        "для генерации высококачественного контента для дипломных работ.\n\n"
        "Мы стремимся облегчить процесс написания дипломных работ, "
        "предоставляя структурированную информацию и идеи для дальнейшего развития. "
    )
    await message.reply(about_text)

async def periodic_save():
    while True:
        await asyncio.sleep(300)  # Сохраняем каждые 5 минут
        save_data(user_data, USER_DATA_FILE)
        save_data(feedback_storage, FEEDBACK_FILE)
        logger.info("Данные пользователей и отзывы сохранены")

async def main():
    # Инициализация бота и диспетчера
    dp.startup.register(on_startup)
    
    # Запуск периодического сохранения
    asyncio.create_task(periodic_save())
    
    # Запуск поллинга
    await dp.start_polling(bot)

async def on_startup(dispatcher: Dispatcher):
    logger.info("Бот запущен")

if __name__ == '__main__':
    asyncio.run(main())
